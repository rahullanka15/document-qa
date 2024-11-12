import time
import streamlit as st
import google.generativeai as genai
import PyPDF2
import docx
import asyncio
import requests
from concurrent.futures import ThreadPoolExecutor
from typing import Dict, List, Optional, Tuple
import os
import logging
from sentence_transformers import SentenceTransformer, util
import torch
import gc
import io
from PIL import Image
import warnings
import re
from pathlib import Path

# Suppress warnings
warnings.filterwarnings('ignore')

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Page configuration must come first
st.set_page_config(
    page_title="📚 Document Assistant",
    page_icon="📚",
    layout="wide"
)

# API Configuration
def setup_apis():
    """Setup and verify API configurations"""
    if 'GOOGLE_API_KEY' not in st.secrets:
        st.error("Please set GOOGLE_API_KEY in streamlit secrets")
        st.stop()
    if 'HF_API_KEY' not in st.secrets:
        st.error("Please set HF_API_KEY in streamlit secrets")
        st.stop()

    # Configure APIs
    genai.configure(api_key=st.secrets['GOOGLE_API_KEY'])
    
    return {
        'model': genai.GenerativeModel('gemini-1.5-pro'),
        'hf_key': st.secrets["HF_API_KEY"],
        'summary_url': "https://api-inference.huggingface.co/models/facebook/bart-large-cnn",
        #'image_url': "https://api-inference.huggingface.co/models/black-forest-labs/FLUX.1-schnell",
        'image_url': "https://api-inference.huggingface.co/models/stabilityai/stable-diffusion-3.5-large",
    }

# Initialize APIs
API_CONFIG = setup_apis()
HEADERS = {
    "Authorization": f"Bearer {API_CONFIG['hf_key']}",
    "Content-Type": "application/json"
}

# Session state initialization
def init_session_state():
    """Initialize session state variables"""
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'documents' not in st.session_state:
        st.session_state.documents = {}
    if 'active_docs' not in st.session_state:
        st.session_state.active_docs = set()
    if 'processing_status' not in st.session_state:
        st.session_state.processing_status = {}
    if 'previous_files' not in st.session_state:
        st.session_state.previous_files = set()

# Constants
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_BATCH_SIZE = 5
MAX_SUMMARY_LENGTH = 250
MIN_SUMMARY_LENGTH = 50

class TextProcessor:
    """Utility class for text processing operations"""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize text content"""
        if not text:
            return ""
        
        # Basic cleaning
        text = text.replace('\n', ' ')
        text = text.replace('\t', ' ')
        
        # Remove multiple spaces
        while '  ' in text:
            text = text.replace('  ', ' ')
            
        # Remove special characters but keep essential punctuation
        text = re.sub(r'[^\w\s.,!?-]', '', text)
        
        return text.strip()
    
    @staticmethod
    def extract_title(text: str, filename: str) -> str:
        """Extract title from text or use filename"""
        if not text:
            return filename
            
        # Try to get first meaningful line
        lines = text.split('\n')
        for line in lines:
            cleaned = line.strip()
            if len(cleaned) > 5 and len(cleaned.split()) <= 20:
                return cleaned
                
        return filename
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            chunks.append(chunk)
            
        return chunks
    
    @staticmethod
    def clean_table_text(text: str) -> str:
        """Clean text extracted from tables"""
        # Remove table formatting characters
        text = re.sub(r'[\|\+\-\=]+', ' ', text)
        # Normalize spaces
        text = ' '.join(text.split())
        return text.strip()
    
class PDFProcessor:
    """Handles PDF document processing"""
    
    @staticmethod
    def process(file) -> Dict[str, str]:
        """Process PDF files with enhanced text extraction"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text_content = []
            
            # Process each page
            for page in pdf_reader.pages:
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    # Clean and add to content
                    cleaned_text = TextProcessor.clean_text(page_text)
                    if cleaned_text:
                        text_content.append(cleaned_text)
            
            # Combine all text
            full_text = '\n'.join(text_content)
            
            if not full_text:
                raise ValueError("No text could be extracted from PDF")
                
            return {
                "content": full_text,
                "type": "pdf",
                "name": file.name,
                "num_pages": len(pdf_reader.pages)
            }
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            raise

class DocxProcessor:
    """Handles DOCX document processing"""
    
    @staticmethod
    def process(file) -> Dict[str, str]:
        """Process DOCX files with formatting preservation"""
        try:
            doc = docx.Document(file)
            text_content = []
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Handle different formatting
                    formatted_text = []
                    for run in paragraph.runs:
                        # Preserve text with its formatting
                        formatted_text.append(run.text)
                    
                    clean_text = TextProcessor.clean_text(' '.join(formatted_text))
                    if clean_text:
                        text_content.append(clean_text)
            
            # Process tables
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = TextProcessor.clean_text(cell.text)
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_content.append(' | '.join(row_text))
                if table_content:
                    text_content.append('\n'.join(table_content))
            
            # Combine all content
            full_text = '\n'.join(text_content)
            
            if not full_text:
                raise ValueError("No text could be extracted from DOCX")
                
            return {
                "content": full_text,
                "type": "docx",
                "name": file.name,
                "has_tables": len(doc.tables) > 0
            }
        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            raise

class TxtProcessor:
    """Handles TXT document processing"""
    
    @staticmethod
    def process(file) -> Dict[str, str]:
        """Process TXT files with encoding handling"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    text = file.read().decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("Could not decode text file with supported encodings")
            
            # Clean text
            clean_text = TextProcessor.clean_text(text)
            
            if not clean_text:
                raise ValueError("No valid text content in file")
                
            return {
                "content": clean_text,
                "type": "txt",
                "name": file.name
            }
        except Exception as e:
            logger.error(f"TXT processing error: {str(e)}")
            raise

# After your imports and configurations, before DocumentProcessor class

def get_combined_context(active_docs: List[str]) -> str:
    """Combine content from multiple documents"""
    combined_text = ""
    for doc_name in active_docs:
        doc = st.session_state.documents.get(doc_name)
        if doc:
            combined_text += f"\nDocument: {doc_name}\n"
            if 'summary' in doc:
                combined_text += f"Summary: {doc['summary']}\n"
            content_preview = doc['content'][:1000]
            combined_text += f"Content Preview: {content_preview}\n"
            combined_text += "---\n"
    return combined_text

def get_friendly_response(question: str, active_docs: List[str]) -> str:
    """Generate response using Gemini model"""
    try:
        context = get_combined_context(active_docs)
        prompt = f"""You are a helpful assistant analyzing these documents. 
        Here are the documents and their summaries:
        {context}

        User question: {question}

        Please provide a comprehensive answer based on ALL selected documents. 
        If referring to specific information, mention which document it came from.
        If you can't find the information in any document, say so clearly.
        Maintain a friendly, professional, conversational tone and explain in simple easy terms.
        
        Guidelines:
        - Cite specific documents when referencing information
        - Be clear about uncertain or missing information
        - Use examples from the documents when relevant
        - Keep the response focused and concise
        """

        response = API_CONFIG['model'].generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.7,
                max_output_tokens=1024,
            )
        )
        return response.text
        
    except Exception as e:
        logger.error(f"Response generation error: {str(e)}")
        return f"I apologize, but I encountered an error: {str(e)}"

def get_similarity_interpretation(score: float) -> Dict[str, str]:
    """Get interpretation of similarity score"""
    if score > 70:
        return {
            "level": "High",
            "description": "Documents are very similar in content and context",
            "color": "green"
        }
    elif score > 40:
        return {
            "level": "Moderate",
            "description": "Documents share some common elements",
            "color": "orange"
        }
    else:
        return {
            "level": "Low",
            "description": "Documents are substantially different",
            "color": "red"
        }

class DocumentProcessor:
    """Main document processing class"""
    
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=5)
        self.processors = {
            'pdf': PDFProcessor.process,
            'docx': DocxProcessor.process,
            'txt': TxtProcessor.process
        }

    async def process_single_file(self, file) -> Optional[Dict]:
        """Process a single file with size and format validation"""
        try:
            # Check file size
            if file.size > MAX_FILE_SIZE:
                raise ValueError(
                    f"File too large: {file.name} "
                    f"({file.size/1024/1024:.1f}MB). Maximum size is "
                    f"{MAX_FILE_SIZE/1024/1024}MB"
                )

            # Validate format
            file_ext = Path(file.name).suffix[1:].lower()
            if file_ext not in self.processors:
                raise ValueError(f"Unsupported file format: {file_ext}")

            # Process document
            result = await asyncio.get_event_loop().run_in_executor(
                self.executor,
                self.processors[file_ext],
                file
            )

            if result and result.get('content'):
                # Add metadata
                result['stats'] = {
                    'word_count': len(result['content'].split()),
                    'char_count': len(result['content']),
                    'upload_time': time.time()
                }
                return result
            
            return None

        except Exception as e:
            logger.error(f"File processing error: {str(e)}")
            st.session_state.processing_status[file.name] = f"Error: {str(e)}"
            return None

    async def get_summary(self, text: str) -> str:
        """Generate summary with chunking and retry logic"""
        try:
            text = TextProcessor.clean_text(text)
            if not text:
                return "No valid text content to summarize"

            # Handle different text lengths
            chunks = TextProcessor.chunk_text(text)
            summaries = []

            for chunk in chunks:
                try:
                    payload = {
                        "inputs": chunk,
                        "parameters": {
                            "max_length": MAX_SUMMARY_LENGTH,
                            "min_length": MIN_SUMMARY_LENGTH,
                            "do_sample": False
                        }
                    }

                    response = requests.post(
                        API_CONFIG['summary_url'],
                        headers=HEADERS,
                        json=payload,
                        timeout=30
                    )

                    if response.status_code == 200:
                        summaries.append(response.json()[0]['summary_text'])
                    elif response.status_code == 503:
                        await asyncio.sleep(20)  # Wait for model to load
                        response = requests.post(
                            API_CONFIG['summary_url'],
                            headers=HEADERS,
                            json=payload
                        )
                        if response.status_code == 200:
                            summaries.append(response.json()[0]['summary_text'])

                    await asyncio.sleep(2)  # Rate limit handling

                except Exception as e:
                    logger.error(f"Chunk summary error: {str(e)}")
                    continue

            if summaries:
                if len(summaries) == 1:
                    return summaries[0]
                else:
                    combined = " ".join(summaries)
                    if len(combined.split()) > MAX_SUMMARY_LENGTH:
                        # Recursively summarize the combined summaries
                        return await self.get_summary(combined)
                    return combined + "\n(Note: Combined from multiple sections)"

            return "Could not generate summary. Please try with simpler content."

        except Exception as e:
            logger.error(f"Summarization error: {str(e)}")
            return f"Error generating summary: {str(e)}"

    async def generate_image(self, text: str, title: str = "") -> Tuple[Optional[Image.Image], Optional[str]]:
        """Generate image with enhanced prompt engineering"""
        try:
            if not text:
                logger.error("No text provided for image generation")
                return None, "No text provided for image generation"

            # Create enhanced prompt
            prompt = f"""Create a visualization about:
            Title: {title}
            Content Summary: {text}
            Style Requirements:
            - Professional and clear visualization
            - Easy to understand representation
            - Focus on key concepts and themes
            - Maintain academic/professional style
            - No words. Focus on Visual learning.
            """

            payload = {"inputs": prompt}

            try:
                response = requests.post(
                    API_CONFIG['image_url'],
                    headers=HEADERS,
                    json=payload,
                    timeout=30
                )

                if response.status_code == 200:
                    image = Image.open(io.BytesIO(response.content))
                    return image, None
                else:
                    error_msg = response.json().get('error', 'Unknown error')
                    if "Max requests total reached" in error_msg:
                        return None, "⏳ Rate limit reached. Please wait 60 seconds..."
                    else:
                        logger.error(f"Image generation failed: {error_msg}")
                        return None, f"Failed to generate image: {error_msg}"

            except requests.exceptions.RequestException as e:
                return None, f"Request failed: {str(e)}"

        except Exception as e:
            return None, f"Error generating image: {str(e)}"

    async def process_files_parallel(self, files: List) -> None:
        """Process multiple files in parallel with batch limiting"""
        start_time = time.time()

        if len(files) > MAX_BATCH_SIZE:
            st.warning(f"Processing files in batches of {MAX_BATCH_SIZE}...")

        tasks = []
        for file in files:
            if file.name not in st.session_state.documents:
                st.session_state.processing_status[file.name] = "Processing..."
                task = asyncio.create_task(self.process_single_file(file))
                tasks.append(task)

        for completed_task in asyncio.as_completed(tasks):
            result = await completed_task
            if result:
                file_name = result['name']
                st.session_state.documents[file_name] = result
                st.session_state.active_docs.add(file_name)
                st.session_state.processing_status[file_name] = "Completed"
                st.success(f"✅ {file_name} processed successfully!")

        processing_time = time.time() - start_time
        logger.info(f"Total processing time: {processing_time:.2f} seconds")
        gc.collect()

class DocumentSimilarity:
    """Handles document similarity calculations"""
    
    def __init__(self):
        try:
            with st.spinner("Loading similarity model..."):
                device = 'cuda' if torch.cuda.is_available() else 'cpu'
                self.model = SentenceTransformer(
                    #'sentence-transformers/all-MiniLM-L6-v2',
                    'sentence-transformers/all-mpnet-base-v2',
                    device=device
                )
        except Exception as e:
            st.error(f"Error initializing similarity model: {str(e)}")
            logger.error(f"Model initialization error: {str(e)}")
            self.model = None

    def calculate_similarity(self, source_text: str, comparison_texts: List[str]) -> Optional[List[float]]:
        """Calculate similarity scores between documents"""
        try:
            if self.model is None:
                return None

            # Convert texts to tensors
            with torch.no_grad():
                source_embedding = self.model.encode(
                    source_text,
                    convert_to_tensor=True
                )
                comparison_embeddings = self.model.encode(
                    comparison_texts,
                    convert_to_tensor=True
                )
                
                # Calculate cosine similarity
                similarities = util.pytorch_cos_sim(
                    source_embedding,
                    comparison_embeddings
                )[0]
            
            # Convert to percentages
            return [float(score) * 100 for score in similarities]

        except Exception as e:
            logger.error(f"Similarity calculation error: {str(e)}")
            st.error(f"Error calculating similarity: {str(e)}")
            return None

    @staticmethod
    def get_similarity_color(score: float) -> str:
        """Get color coding for similarity score"""
        if score > 70:
            return 'green'
        elif score > 40:
            return 'orange'
        else:
            return 'red'
        
class UIComponents:
    """Handles all UI components and layouts"""
    
    # @staticmethod
    # def render_chat_interface(container):
    #     """Render chat interface with message history"""
    #     for message in st.session_state.chat_history:
    #         with container.chat_message(message["role"]):
    #             st.markdown(message["content"])

    @staticmethod
    def render_document_stats(doc: Dict):
        """Render document statistics"""
        st.markdown(f"""
        **Document Statistics:**
        - Words: {len(doc['content'].split())}
        - Characters: {len(doc['content'])}
        - Type: {doc['type'].upper()}
        {f"- Pages: {doc['num_pages']}" if 'num_pages' in doc else ""}
        """)

    @staticmethod
    def render_document_content(doc: Dict, key_prefix: str):
        """Render document content and summary"""
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### Original Content")
            st.write("")
            st.text_area(
                label="",
                value=doc['content'][:1000] + "..." if len(doc['content']) > 1000 else doc['content'],
                height=300,
                disabled=True,
                key=f"content_{key_prefix}"
            )
        
        with col2:
            st.markdown("### Summary")
            st.write("")
            if 'summary' in doc:
                st.text_area(
                    label="",
                    value=doc['summary'],
                    height=300,
                    disabled=True,
                    key=f"summary_{key_prefix}"
                )
            else:
                st.info("Click 'Generate Summaries' to create summary")

    @staticmethod
    async def handle_rate_limit():
        """Handle API rate limit with countdown"""
        time_left = 60
        progress_text = "Please wait..."
        progress_bar = st.progress(0)
        
        for i in range(time_left):
            progress_bar.progress((i + 1) / time_left)
            st.text(f"Time remaining: {time_left - i} seconds")
            await asyncio.sleep(1)
            
        st.success("You can try again now!")
        st.rerun()
    
def render_similarity_tab():
    """Render similarity analysis interface"""
    st.markdown("### Document Similarity Analysis")
    
    if len(st.session_state.active_docs) < 2:
        st.warning("Please select at least 2 documents to use similarity analysis.")
        return
    
    available_docs = list(st.session_state.active_docs)
    
    selected_doc = st.selectbox(
        "Select a document to compare with others",
        available_docs,
        key="similarity_doc_selector"
    )
    
    if selected_doc:
        source_doc = st.session_state.documents[selected_doc]
        
        comparison_docs = {
            name: doc for name, doc in st.session_state.documents.items() 
            if name != selected_doc and name in st.session_state.active_docs
        }
        
        if comparison_docs:
            with st.spinner("Calculating similarity scores..."):
                try:
                    similarity_calc = DocumentSimilarity()
                    
                    comparison_texts = [doc['content'] for doc in comparison_docs.values()]
                    scores = similarity_calc.calculate_similarity(
                        source_doc['content'],
                        comparison_texts
                    )
                    
                    if scores:
                        st.markdown("### Similarity Scores")
                        
                        for (doc_name, doc), score in zip(comparison_docs.items(), scores):
                            with st.expander(f"**{doc_name}**", expanded=True):
                                col1, col2 = st.columns([3, 1])
                                
                                with col1:
                                    if 'summary' in doc:
                                        st.markdown("**Summary:**")
                                        st.info(doc['summary'][:200] + "...")
                                
                                with col2:
                                    color = DocumentSimilarity.get_similarity_color(score)
                                    st.markdown(
                                        f"<h3 style='text-align: center; color: {color}'>"
                                        f"{score:.1f}%</h3>",
                                        unsafe_allow_html=True
                                    )
                                    
                                    interp = get_similarity_interpretation(score)
                                    if score > 70:
                                        st.success(interp["description"])
                                    elif score > 40:
                                        st.warning(interp["description"])
                                    else:
                                        st.error(interp["description"])
                                
                                st.markdown("---")
                    else:
                        st.error("Error calculating similarity scores")
                
                except Exception as e:
                    st.error(f"Error in similarity analysis: {str(e)}")
                    logger.error(f"Similarity error details: {str(e)}")

class DocumentTabs:
    """Handles different tab views in the application"""

    def __init__(self, doc_processor: DocumentProcessor):
        self.doc_processor = doc_processor

    def render_chat_tab(self, tab):
        """Render chat interface tab with fixed input and proper message flow"""
        with tab:
            st.markdown("### Chat with your Documents")
            
            if st.session_state.active_docs:
                st.info(f"📚 Currently analyzing: {', '.join(st.session_state.active_docs)}")
                
                # Create containers for messages and input
                messages_container = st.container()
                input_container = st.container()
                
                # Handle input first (at the bottom)
                with input_container:
                    # Add some spacing before the input
                    st.markdown("<br>" * 2, unsafe_allow_html=True)
                    
                    # Get user input
                    prompt = st.chat_input("Ask me anything about your documents...")
                
                # Display messages in the messages container
                with messages_container:
                    for message in st.session_state.chat_history:
                        with st.chat_message(message["role"]):
                            st.markdown(message["content"])
                    
                    # Handle new message if there's input
                    if prompt:
                        # Immediately show user message
                        with st.chat_message("user"):
                            st.markdown(prompt)
                        
                        # Show assistant response with loading indicator
                        with st.chat_message("assistant"):
                            with st.spinner("Thinking..."):
                                response = get_friendly_response(
                                    prompt,
                                    list(st.session_state.active_docs)
                                )
                                st.markdown(response)
                        
                        # Update chat history
                        st.session_state.chat_history.extend([
                            {"role": "user", "content": prompt},
                            {"role": "assistant", "content": response}
                        ])
                        
                        # Rerun to update the display properly
                        st.rerun()
            else:
                st.info("👈 Please upload and select documents to start chatting!")


    def render_documents_tab(self, tab):
        """Render documents analysis tab"""
        with tab:
            st.markdown("### Document Summaries")
            
            # Document selector with default to all documents
            available_docs = list(st.session_state.documents.keys())
            selected_docs = st.multiselect(
                "Select documents to summarize",
                available_docs,
                default=available_docs,  # Default to all documents
                key="summary_doc_selector"
            )
            
            if selected_docs:
                if st.button("Generate Summaries", key="generate_summaries"):
                    total_docs = len(selected_docs)
                    progress_bar = st.progress(0)
                    
                    for idx, doc_name in enumerate(selected_docs):
                        if doc_name in st.session_state.documents:
                            doc = st.session_state.documents[doc_name]
                            status_text = st.empty()
                            
                            if 'summary' not in doc:
                                try:
                                    status_text.text(f"Processing {doc_name}...")
                                    summary = asyncio.run(self.doc_processor.get_summary(doc['content']))
                                    if summary:
                                        st.session_state.documents[doc_name]['summary'] = summary
                                        st.success(f"✅ Summary generated for {doc_name}")
                                except Exception as e:
                                    st.error(f"Error generating summary for {doc_name}: {str(e)}")
                            else:
                                st.info(f"Summary already exists for {doc_name}")
                            
                            # Update progress
                            progress_bar.progress((idx + 1) / total_docs)
                    
                    st.success("All documents processed!")
            
            # Display documents
            for doc_name in selected_docs:
                if doc_name in st.session_state.documents:
                    doc = st.session_state.documents[doc_name]
                    with st.expander(f"📄 {doc_name}", expanded=True):
                        UIComponents.render_document_content(doc, doc_name)
                        UIComponents.render_document_stats(doc)

    def render_image_tab(self, tab):
        """Render image generation tab"""
        with tab:
            st.markdown("### Document Visualization")
            
            # Document selector
            available_docs = list(st.session_state.documents.keys())
            selected_doc = st.selectbox(
                "Select a document to visualize",
                [""] + available_docs,
                key="image_doc_selector"
            )
            
            if selected_doc:
                doc = st.session_state.documents.get(selected_doc)
                if doc:
                    col1, col2 = st.columns([2, 1])
                    with col1:
                        generate_btn = st.button(
                            "Generate Image",
                            key="generate_image",
                            use_container_width=True
                        )
                    
                    if generate_btn:
                        asyncio.run(self.handle_image_generation(doc, selected_doc))
                    
                    # Display existing image
                    if 'image' in doc and doc['image'] is not None:
                        st.markdown("### Generated Image")
                        st.image(doc['image'], use_container_width=True)
                        
                        if st.button("🔄 Generate New Image", key="regenerate"):
                            asyncio.run(self.handle_image_generation(doc, selected_doc))

    async def handle_image_generation(self, doc: Dict, doc_name: str):
        """Handle image generation process"""
        with st.spinner(f"Processing {doc_name}..."):
            try:
                # Ensure summary exists
                if 'summary' not in doc or not doc['summary']:
                    st.info("Generating summary first...")
                    summary = await self.doc_processor.get_summary(doc['content'])
                    if summary:
                        st.session_state.documents[doc_name]['summary'] = summary
                        st.success("Summary generated!")
                    else:
                        st.error("Failed to generate summary")
                        return

                # Extract title
                title = TextProcessor.extract_title(doc['content'], doc_name)
                
                # Generate image
                st.info("Generating image...")
                image, error_msg = await self.doc_processor.generate_image(
                    doc['summary'],
                    title=title
                )
                
                if image:
                    st.session_state.documents[doc_name]['image'] = image
                    st.success("✅ Image generated successfully!")
                    st.rerun()
                elif "Rate limit" in str(error_msg):
                    await UIComponents.handle_rate_limit()
                else:
                    st.error(error_msg)
                    
            except Exception as e:
                st.error(f"Error: {str(e)}")
                logger.error(f"Error details: {str(e)}")

def main():
    """Main application entry point"""
    # Initialize session state
    init_session_state()
    
    # Setup document processor
    doc_processor = DocumentProcessor()
    tabs_handler = DocumentTabs(doc_processor)
    
    # Main UI
    st.title("📚 Document Assistant")
    st.markdown("Upload your documents for summaries and interactive chat!")
    
    # Sidebar
    with st.sidebar:
        render_sidebar(doc_processor)
    
    # Main content
    if st.session_state.active_docs:
        render_main_content(tabs_handler)
    else:
        st.info("👈 Please upload and select documents to get started!")
    
    # Footer
    st.markdown("---")
    st.markdown("💡 Powered by Gemini & BART | Made with Streamlit")

def render_sidebar(doc_processor: DocumentProcessor):
    """Render sidebar content"""
    st.header("📎 Document Management")
    
    # File uploader
    uploaded_files = st.file_uploader(
        "Upload documents",
        type=['pdf', 'docx', 'txt'],
        accept_multiple_files=True,
        key="file_uploader"
    )
    
    # Handle file tracking and removal
    current_files = set(f.name for f in uploaded_files) if uploaded_files else set()
    deleted_files = st.session_state.previous_files - current_files
    
    # Process deleted files
    for deleted_file in deleted_files:
        if deleted_file in st.session_state.documents:
            del st.session_state.documents[deleted_file]
            st.session_state.active_docs.discard(deleted_file)
            if deleted_file in st.session_state.processing_status:
                del st.session_state.processing_status[deleted_file]
    
    # Update tracked files
    st.session_state.previous_files = current_files
    
    # Process new files
    if uploaded_files:
        asyncio.run(doc_processor.process_files_parallel(uploaded_files))
    
    # Document selection
    if st.session_state.documents:
        st.markdown("### Selected Documents")
        
        if st.button("Deselect All"):
            st.session_state.active_docs = set()
            st.rerun()
        
        for doc_name in st.session_state.documents:
            checkbox = st.checkbox(
                f"📄 {doc_name}",
                key=f"check_{doc_name}",
                value=doc_name in st.session_state.active_docs
            )
            
            if checkbox:
                st.session_state.active_docs.add(doc_name)
            else:
                st.session_state.active_docs.discard(doc_name)

def render_main_content(tabs_handler: DocumentTabs):
    """Render main content area with tabs"""
    # Determine which tabs to show
    show_similarity = len(st.session_state.active_docs) >= 2
    tabs = ["💭 Chat", "📑 Documents", "🎨 Images"]
    if show_similarity:
        tabs.append("🔄 Similarity")
    
    # Create tabs
    selected_tabs = st.tabs(tabs)
    
    # Render each tab
    tabs_handler.render_chat_tab(selected_tabs[0])
    tabs_handler.render_documents_tab(selected_tabs[1])
    tabs_handler.render_image_tab(selected_tabs[2])
    
    if show_similarity:
        with selected_tabs[3]:
            render_similarity_tab()

if __name__ == "__main__":
    main()