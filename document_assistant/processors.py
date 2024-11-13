# processors.py
import re
from document_assistant.core import logger, MAX_FILE_SIZE
from pathlib import Path
from typing import Dict, List, Optional
import PyPDF2
import docx
import tempfile
import os
import subprocess

class PDFProcessor:
    """Handles PDF document processing"""
    
    @staticmethod
    def process(file) -> Dict[str, str]:
        """Process PDF files with enhanced text extraction"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text_content = []
            
            # Process each page
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    cleaned_text = TextProcessor.clean_text(page_text)
                    if cleaned_text:
                        text_content.append(cleaned_text)
            
            # Combine all text
            full_text = '\n'.join(text_content)
            
            if not full_text:
                raise ValueError("No text could be extracted from PDF")
            
            # Calculate statistics
            stats = {
                "type": "pdf",
                "name": file.name,
                "file_size": file.size,
                "num_pages": len(pdf_reader.pages),
                "word_count": len(full_text.split()),
                "char_count": len(full_text)
            }
                
            return {
                "content": full_text,
                "stats": stats,
                "raw_text": full_text,  # Keep original for other processors
                "processed": True
            }
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            raise

class DocxProcessor:
    """Handles DOCX document processing"""
    
    @staticmethod
    def process(file) -> Dict[str, str]:
        try:
            doc = docx.Document(file)
            text_content = []
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    formatted_text = [run.text for run in paragraph.runs]
                    clean_text = TextProcessor.clean_text(' '.join(formatted_text))
                    if clean_text:
                        text_content.append(clean_text)
            
            # Process tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                table_content = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = TextProcessor.clean_text(cell.text)
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_content.append(' | '.join(row_text))
                if table_content:
                    text_content.append('\n'.join(table_content))
            
            # Combine all content
            full_text = '\n'.join(text_content)
            
            if not full_text:
                raise ValueError("No text could be extracted from DOCX")
            
            # Calculate statistics
            stats = {
                "type": "docx",
                "name": file.name,
                "file_size": file.size,
                "table_count": table_count,
                "paragraph_count": len(doc.paragraphs),
                "word_count": len(full_text.split()),
                "char_count": len(full_text)
            }
                
            return {
                "content": full_text,
                "stats": stats,
                "raw_text": full_text,
                "processed": True
            }
        except Exception as e:
            logger.error(f"DOCX processing error: {str(e)}")
            raise

class DocProcessor:
    """Handles DOC document processing"""
    
    @staticmethod
    def process(file) -> Dict[str, str]:
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_file:
                temp_file.write(file.read())
                temp_path = temp_file.name
            
            try:
                # Try antiword first
                text = subprocess.check_output(['antiword', temp_path]).decode('utf-8')
            except (subprocess.SubprocessError, FileNotFoundError):
                try:
                    # Fallback to python-docx
                    doc = docx.Document(temp_path)
                    text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
                except Exception as e:
                    raise ValueError(f"Could not process DOC file: {str(e)}")
            finally:
                os.unlink(temp_path)
            
            clean_text = TextProcessor.clean_text(text)
            
            if not clean_text:
                raise ValueError("No valid text content in DOC file")
            
            # Calculate statistics
            stats = {
                "type": "doc",
                "name": file.name,
                "file_size": file.size,
                "word_count": len(clean_text.split()),
                "char_count": len(clean_text)
            }
                
            return {
                "content": clean_text,
                "stats": stats,
                "raw_text": text,
                "processed": True
            }
            
        except Exception as e:
            logger.error(f"DOC processing error: {str(e)}")
            raise

class TxtProcessor:
    """Handles TXT document processing"""
    
    @staticmethod
    def process(file) -> Dict[str, str]:
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    text = file.read().decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("Could not decode text file with supported encodings")
            
            clean_text = TextProcessor.clean_text(text)
            
            if not clean_text:
                raise ValueError("No valid text content in file")
            
            # Calculate statistics
            stats = {
                "type": "txt",
                "name": file.name,
                "file_size": file.size,
                "line_count": len(text.splitlines()),
                "word_count": len(clean_text.split()),
                "char_count": len(clean_text)
            }
                
            return {
                "content": clean_text,
                "stats": stats,
                "raw_text": text,
                "processed": True
            }
        except Exception as e:
            logger.error(f"TXT processing error: {str(e)}")
            raise

class TextProcessor:
    """Utility class for text processing operations"""
    
    @staticmethod
    def clean_text(text: str) -> str:
        if not text:
            return ""
        
        # Basic cleaning
        text = text.replace('\n', ' ').replace('\t', ' ')
        
        # Remove multiple spaces
        while '  ' in text:
            text = text.replace('  ', ' ')
            
        # Remove special characters but keep essential punctuation
        text = re.sub(r'[^\w\s.,!?-]', '', text)
        
        return text.strip()
    
    @staticmethod
    def extract_title(text: str, filename: str) -> str:
        if not text:
            return filename
            
        lines = text.split('\n')
        for line in lines:
            cleaned = line.strip()
            if len(cleaned) > 5 and len(cleaned.split()) <= 20:
                return cleaned
                
        return filename
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            chunks.append(chunk)
            
        return chunks